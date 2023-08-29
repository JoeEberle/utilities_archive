#!/usr/bin/env python
# coding: utf-8
# In[3]:

__author__      = "Joe Eberle"
__copyright__   = " "


def get_solution_documentation():
    ''' the get_solution_documentation function will return the solution documentation in text form. '''  
    solution_documentation = '\n'
    solution_documentation = solution_documentation + get_solution_introduction()
    solution_documentation = solution_documentation + get_credits() 
    solution_documentation = solution_documentation + get_solution_steps()
    solution_documentation = solution_documentation + get_solution_notes()
#     solution_documentation = solution_documentation + get_magic_or_dunder()    
    return solution_documentation  

def create_solution_documentation_file(documentation_file_name = 'C://working_directory//documentation//Test_Document.doc'):
    ''' the get_solution_documentation function will create a formatted word document. ''' 
    file_name = create_solution_documentation_file(documentation_file_name) 
    return file_name  

def save_documentation(filename):
    ''' the save_documentation function will save the solution documentation in a file.'''  
    documentation = get_solution_documentation()      # get all of the documentation 
    f = open(filename, "w")
    f.write(documentation)
    f.close()
     
def get_magic_or_dunder(object):
    ''' the get_magic_or_dunder function will return all of the python magic functions for the class. '''      
    magic_or_dunder = '\n'  # Start with a new line to separate documentation and make it more readable.   
    magic_or_dunder = magic_or_dunder + 'Magic or Dunder:\n'  
    magic_or_dunder = magic_or_dunder + '\n \n __builtins__:' + object.__builtins__ 
    magic_or_dunder = magic_or_dunder + '\n \n __cached__:' + object.__cached__ 
    magic_or_dunder = magic_or_dunder + '\n \n __doc__:' + object.__doc__ 
    magic_or_dunder = magic_or_dunder + '\n \n __file__:' + object.__file__ 
    magic_or_dunder = magic_or_dunder + '\n \n __loader__:' + object.__loader__ 
    magic_or_dunder = magic_or_dunder + '\n \n __name__:' + object.__name__ 
    magic_or_dunder = magic_or_dunder + '\n \n __package__:' + object.__package__ 
    magic_or_dunder = magic_or_dunder + '\n \n __spec__:' + object.__spec__ 
    magic_or_dunder = magic_or_dunder + '\n \n __spec__:' + object.__author__     
    return magic_or_dunder

# get_credits - returns who contributed to the solution and when.
def get_credits(): 
    ''' the get_credits function will return who contributed to the solution and when. '''         
    solution_credits = '\n\n'   # Start with a new line to separate documentation and make it more readable.   
    solution_credits = solution_credits + '\nThis data science solution was developed in collaboration with Joe Eberle.' 
    solution_credits = solution_credits + '\nThe solution was developed in Python starting on 11/13/2022'
    solution_credits = solution_credits + '\nThis solution is free and open source and the code is openly available for general use.'   
    return solution_credits 
    
# get_solution_steps- Overview of solution Process Steps 
def get_solution_steps():
    ''' the get_solution_steps function will return high levels steps needed for the solution. '''        
    solution_steps = '\n'   # Start with a new line to separate documentation and make it more readable.     
    solution_steps = solution_steps + 'The steps for this solution are:\n'
    solution_steps = solution_steps + 'Step 1: identify the directory to format excel files.\n'
    solution_steps = solution_steps + 'Step 2: run the remote process to format the excel files .\n'    
    solution_steps = solution_steps + 'Step 3: optional test - Visually inspect a sample of  your modified excel files.\n'      
    return solution_steps 
    
# Introduction - Overview of Solution  
def get_solution_introduction():
    ''' the get_solution_introduction function returns the and overview or introduction of what the solution does. '''  
    solution_introduction = '\n'   # Start with a new line to separate documentation and make it more readable.     
    solution_introduction = solution_introduction + 'This data science solution will enhance and improve your excel files automatically.\n'
    solution_introduction = solution_introduction + 'This helps give all your excel files the same look and feel.\n'   
    solution_introduction = solution_introduction + 'You can easily modify the look and feel of your excel reports to make them your own.\n'      
    solution_introduction = solution_introduction + 'This solution will help you document by adding meta tags to your existing excel.\n'
    solution_introduction = solution_introduction + 'This solution will help eliminate hours of manual formatting.\n'    
    return solution_introduction 

#  get_solution_notes
def get_solution_notes():
    ''' the get_solution_notes function returns any special or helpful notes about the solution.'''  
    solution_notes = '\n \n'   # Start with a new line to separate documentation and make it more readable. 
    solution_notes = solution_notes + 'The additional notes for this solution are: \n'
    solution_notes = solution_notes + 'Note 1: This solution makes and saves a formatted copy of your excel files.\n'  
    solution_notes = solution_notes + 'Note 2: The copy of your excel will have the same name with a _ft added.\n' 
    solution_notes = solution_notes + 'Note 3: This solution uses the popular openpyxl library to improve your spreadsheets .\n'     
    return solution_notes 

def create_solution_documentation_file(filename, solution_name):
    ''' the get_solution_documentation function will create a formatted word document. ''' 
    from docx import Document
    document = Document()
    
    document.add_heading(f'Solution Documentation -  Solution Name {solution_name}', 0)
#     document.add_picture('C://working_directory//images//configuring_settings.png')    
    document.add_page_break()
    document.add_heading('Solution Documentation -  Solution Name ', 0)
    document.add_heading('Solution Introduction', 1)    
    p_intro = document.add_paragraph(get_solution_introduction()) 
    document.add_heading('Solution Credits', 1)    
    p_credits = document.add_paragraph(get_credits())  
    document.add_heading('Solution Steps', 1)    
    p_steps = document.add_paragraph(get_solution_steps())  
    document.add_heading('Solution Notes', 1)    
    p_notes = document.add_paragraph(get_solution_notes())           
    
    document.save(filename)
    return filename

def get_solution_image():
    solution_image = 'C://working_directory//images//' + 'configuring_settings.png'
    return solution_image 

 



