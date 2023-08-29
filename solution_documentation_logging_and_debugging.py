#!/usr/bin/env python
# coding: utf-8
# In[3]:

__author__      = "Joe Eberle"
__copyright__   = " "


def get_solution_documentation():
    ''' the get_solution_documentation function will return the solution documentation in text form. '''  
    solution_documentation = '\n'
    solution_documentation = solution_documentation + get_solution_introduction()
    solution_documentation = solution_documentation + get_credits() 
    solution_documentation = solution_documentation + get_solution_steps()
    solution_documentation = solution_documentation + get_solution_notes()
#     solution_documentation = solution_documentation + get_magic_or_dunder()    
    return solution_documentation  

def create_solution_documentation_file(documentation_file_name):
    ''' the get_solution_documentation function will create a formatted word document. ''' 
    documentation_file_name = 'C://working_directory//documentation//Test_Document.doc'
    file_name = create_solution_documentation_file(documentation_file_name) 
    return file_name  

def save_documentation(filename):
    ''' the save_documentation function will save the solution documentation in a file.'''  
    documentation = get_solution_documentation()      # get all of the documentation 
    f = open(filename, "w")
    f.write(documentation)
    f.close()
     
def get_magic_or_dunder(object):
    ''' the get_magic_or_dunder function will return all of the python magic functions for the class. '''      
    magic_or_dunder = '\n'  # Start with a new line to separate documentation and make it more readable.   
    magic_or_dunder = magic_or_dunder + 'Magic or Dunder:\n'  
    magic_or_dunder = magic_or_dunder + '\n \n __builtins__:' + object.__builtins__ 
    magic_or_dunder = magic_or_dunder + '\n \n __cached__:' + object.__cached__ 
    magic_or_dunder = magic_or_dunder + '\n \n __doc__:' + object.__doc__ 
    magic_or_dunder = magic_or_dunder + '\n \n __file__:' + object.__file__ 
    magic_or_dunder = magic_or_dunder + '\n \n __loader__:' + object.__loader__ 
    magic_or_dunder = magic_or_dunder + '\n \n __name__:' + object.__name__ 
    magic_or_dunder = magic_or_dunder + '\n \n __package__:' + object.__package__ 
    magic_or_dunder = magic_or_dunder + '\n \n __spec__:' + object.__spec__ 
    magic_or_dunder = magic_or_dunder + '\n \n __spec__:' + object.__author__     
    return magic_or_dunder

# get_credits - returns who contributed to the solution and when.
def get_credits(): 
    ''' the get_credits function will return who contributed to the solution and when. '''         
    solution_credits = ''   # Start with a new line to separate documentation and make it more readable.   
    solution_credits = solution_credits + '\nThis data science solution was developed:' 
    solution_credits = solution_credits + '\nIn collaboration by Joe Eberle and others.' 
    solution_credits = solution_credits + '\nThe solution was developed in Python starting on 1/15/2023'
    solution_credits = solution_credits + '\nThis solution is free and open source and the code is openly available for general use.'   
    return solution_credits 
    
# get_solution_steps- Overview of solution Process Steps 
def get_solution_steps():
    ''' the get_solution_steps function will return high levels steps needed for the solution. '''        
    solution_steps = '\n'   # Start with a new line to separate documentation and make it more readable.     
    solution_steps = solution_steps + '\nThe high level steps for this solution are:\n'
    solution_steps = solution_steps + 'Step 1: Turn on the logger engine.\n'
    solution_steps = solution_steps + 'Step 2: Insert logging and debugging throughout code\n'
    solution_steps = solution_steps + 'Step 3: Inspect the logging at the end of the solution.\n' 
    solution_steps = solution_steps + 'Optional Step 4: Inspect the historical log to analyze performnce over time.\n'     
    return solution_steps 
    
# Introduction - Overview of Solution  
def get_solution_introduction():
    ''' the get_solution_introduction function returns the and overview or introduction of what the solution does. '''  
    solution_introduction = '\n'   # Start with a new line to separate documentation and make it more readable.     
    solution_introduction = solution_introduction + 'This solution creates a log of information about the solution itself.\n'
    solution_introduction = solution_introduction +  'The log file enables you to quickly and easily debug and track performance of any solution.\n'
    solution_introduction = solution_introduction +  'The file will be easily modifyable and extensible to track information as needed.\n'  
    solution_introduction = solution_introduction +  'The current log file is automatically appended to the historical log so that performance and functionality can be tracked over time.\n'        
    return solution_introduction 

#  get_solution_notes
def get_solution_notes():
    ''' the get_solution_notes function returns any special or helpful notes about the solution.'''  
    solution_notes = '\n'   # Start with a new line to separate documentation and make it more readable. 
    solution_notes = solution_notes + '\nSome additional notes for this solution are:\n'
    solution_notes = solution_notes + 'Note 1: The log file is kept by default in the current working directory.\n'  
    solution_notes = solution_notes + 'Note 2: Logging is designed to persist even if the process fails.\n' 
    solution_notes = solution_notes + 'Note 3: Each solution gets its own log file to keep things brief and simple.\n'     
    solution_notes = solution_notes + 'Note 4: Each log file gets reset every time the solution is run.\n'
    solution_notes = solution_notes + 'Note 5: Each log file is appended to a historical log for anaylsis over time.\n'    
    solution_notes = solution_notes + 'Note 6: There are 5 levels of logging you can use to organize the log file \n'        
    solution_notes = solution_notes + 'Note 7: It is up to you how much logging you think is necessary for your solution\n'       
    return solution_notes 

def create_solution_documentation_file(filename):
    ''' the get_solution_documentation function will create a formatted word document. ''' 
    from docx import Document
    document = Document()
    
    document.add_heading('Solution Documentation -  Solution Name ', 0)
    document.add_picture('C://working_directory//images//configuring_settings.png')    
    document.add_page_break()
    document.add_heading('Solution Documentation -  Solution Name ', 0)
    document.add_heading('Solution Introduction', 1)    
    p_intro = document.add_paragraph(get_solution_introduction()) 
    document.add_heading('Solution Credits', 1)    
    p_credits = document.add_paragraph(get_credits())  
    document.add_heading('Solution Steps', 1)    
    p_steps = document.add_paragraph(get_solution_steps())  
    document.add_heading('Solution Notes', 1)    
    p_notes = document.add_paragraph(get_solution_notes())           
    
    document.save(filename)
    return filename

def get_solution_image():
    solution_image = 'C://working_directory//images//' + 'configuring_settings.png'
    return solution_image 

 



