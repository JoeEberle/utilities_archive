#!/usr/bin/env python
# coding: utf-8
# In[3]:

__author__      = "Joe Eberle"
__copyright__   = " "


def get_solution_documentation():
    ''' the get_solution_documentation function will return the solution documentation in text form. '''  
    solution_documentation = 'This solution documentation: \n'
    solution_documentation = solution_documentation + get_solution_introduction()
    solution_documentation = solution_documentation + get_credits() 
    solution_documentation = solution_documentation + get_solution_steps()
    solution_documentation = solution_documentation + get_solution_notes()
#     solution_documentation = solution_documentation + get_magic_or_dunder()    
    return solution_documentation  

def create_solution_documentation_file(documentation_file_name):
    ''' the get_solution_documentation function will create a formatted word document. ''' 
    documentation_file_name = 'C://working_directory//documentation//Test_Document.doc'
    file_name = create_solution_documentation_file(documentation_file_name) 
    return file_name  

def save_documentation(filename):
    ''' the save_documentation function will save the solution documentation in a file.'''  
    
    documentation = get_solution_documentation()      # get all of the documentation 
    f = open(filename, "w")
    f.write(documentation)
    f.close()
     
def get_magic_or_dunder(object):
    ''' the get_magic_or_dunder function will return all of the python magic functions for the class. '''      
    magic_or_dunder = '\n'  # Start with a new line to separate documentation and make it more readable.   
    magic_or_dunder = magic_or_dunder + 'Magic or Dunder:\n'  
    magic_or_dunder = magic_or_dunder + '\n \n __builtins__:' + object.__builtins__ 
    magic_or_dunder = magic_or_dunder + '\n \n __cached__:' + object.__cached__ 
    magic_or_dunder = magic_or_dunder + '\n \n __doc__:' + object.__doc__ 
    magic_or_dunder = magic_or_dunder + '\n \n __file__:' + object.__file__ 
    magic_or_dunder = magic_or_dunder + '\n \n __loader__:' + object.__loader__ 
    magic_or_dunder = magic_or_dunder + '\n \n __name__:' + object.__name__ 
    magic_or_dunder = magic_or_dunder + '\n \n __package__:' + object.__package__ 
    magic_or_dunder = magic_or_dunder + '\n \n __spec__:' + object.__spec__ 
    magic_or_dunder = magic_or_dunder + '\n \n __spec__:' + object.__author__     
    return magic_or_dunder

# get_credits - returns who contributed to the solution and when.
def get_credits(): 
    ''' the get_credits function will return who contributed to the solution and when. '''         
    solution_credits = '\n\n'   # Start with a new line to separate documentation and make it more readable.   
    solution_credits = solution_credits + '\nThis data science solution was developed:' 
    solution_credits = solution_credits + '\nIn collaboration by Joe Eberle and others.' 
    solution_credits = solution_credits + '\nThe solution was developed in Python starting on 11/03/2022 and revised on 01/11/2023'
    solution_credits = solution_credits + '\nThis solution is free and open source and the code is openly available for general use.'   
    return solution_credits 
    
# get_solution_steps- Overview of solution Process Steps 
def get_solution_steps():
    ''' the get_solution_steps function will return high levels steps needed for the solution. '''        
    solution_steps = '\n'   # Start with a new line to separate documentation and make it more readable.     
    solution_steps = 'The high level steps for this solution are:\n'
    solution_steps = solution_steps + 'Step 1: Establish the configuration parser engine.\n'
    solution_steps = solution_steps + 'Step 2: Establish all the Sections for initial configuration settings\n'
    solution_steps = solution_steps + 'Step 3: Write out name value pairs for all of the configuration settings.\n' 
    solution_steps = solution_steps + 'Step 4: Repeat steps 2 and 3 until all of the sections and configuration settings are complete.\n'  
    solution_steps = solution_steps + 'Step 5: Write out the config.ini file .\n'   
    solution_steps = solution_steps + 'Step 6: Test a singular Config.ini setting.\n'   
    solution_steps = solution_steps + 'Step 7: Print out the entire config.ini file and visually inspect it.\n'       
    return solution_steps 
    
# Introduction - Overview of Solution  
def get_solution_introduction():
    ''' the get_solution_introduction function returns the and overview or introduction of what the solution does. '''  
    solution_introduction = '\n'   # Start with a new line to separate documentation and make it more readable.     
    solution_introduction = solution_introduction + 'This data science solution will establish a configuration file that contains configuration settings.\n'
    solution_introduction = solution_introduction +  'The configuration file will allow you to quickly and easily customize settings and global variables so that solutions are more easily modified.\n'
    solution_introduction = solution_introduction +  'The configuration file will be easily modifyable allowing users to reset configurations as needed.\n'  
    solution_introduction = solution_introduction +  'Using configuration file allows the developer to avoid hard coding and to obfuscate any sensitive settings usch as passwords, etc.\n'        
    return solution_introduction 

#  get_solution_notes
def get_solution_notes():
    ''' the get_solution_notes function returns any special or helpful notes about the solution.'''  
    solution_notes = '\n \n'   # Start with a new line to separate documentation and make it more readable. 
    solution_notes = solution_notes + 'The additional notes for this solution are: \n'
    solution_notes = solution_notes + 'Note 1: The configuration file name config.ini is the default file name is located by default in the current working directory.\n'  
    solution_notes = solution_notes + 'Note 2: The global_infrastructure logging setting is set(True) as the default setting so EVERY solution has event logging turned on for debugging and performance monitoring.\n' 
    solution_notes = solution_notes + 'Note 3: The global_infrastructure documentation setting is set(True) as the default setting so EVERY solution automatically produces its own documentation.\n'       
    solution_notes = solution_notes + 'Note 4: The global_infrastructure section of the configuration file can be overwritten with individual solution settings... however this is the individual developers responsibility to overwrite global settings and manage their own custom settings.\n'
    solution_notes = solution_notes + 'Note 5: A monolithic configuration file approach was chosen as the default approach for this solution meaning that one configuration file is used for ALL solutions, however, there is nothing preventing you from createing and using individual configuration files for each solution you develop. This is personal preference and an architecture decision for you to consider when choosing whats best for you or your organization. \n'    
    solution_notes = solution_notes + 'Note 6: An advantage of using a monolithic configuration file approach is that it promotes standardization and consistency throught the organization. This means that all of your applications can have the same look and feel and consistent behaviour. \n'        
    solution_notes = solution_notes + 'Note 7: This data science solution was developed in python using jupyter notebook for ease of illustration an educational purposes. There is nothing preventing you from using your preferred language and integrated development environment (IDE) of your choice such as visual studio or pycharm. This solution was devloped on a windows platform and has not been tested on linux or othewr operating systems, however it was designed with cross-platform compatible libraries so hypothetically the solution should be portable to other platforms. \n'       
    return solution_notes 



def create_solution_documentation_file(filename):
    ''' the get_solution_documentation function will create a formatted word document. ''' 
    from docx import Document
    document = Document()
    
    document.add_heading('Solution Documentation -  Solution Name ', 0)
    document.add_picture('C://working_directory//images//configuring_settings.png')    
    document.add_page_break()
    document.add_heading('Solution Documentation -  Solution Name ', 0)
    document.add_heading('Solution Introduction', 1)    
    p_intro = document.add_paragraph(get_solution_introduction()) 
    document.add_heading('Solution Credits', 1)    
    p_credits = document.add_paragraph(get_credits())  
    document.add_heading('Solution Steps', 1)    
    p_steps = document.add_paragraph(get_solution_steps())  
    document.add_heading('Solution Notes', 1)    
    p_notes = document.add_paragraph(get_solution_notes())           
    
    document.save(filename)
    return filename

def get_solution_image():
    solution_image = 'C://working_directory//images//' + 'configuring_settings.png'
    return solution_image 

 



